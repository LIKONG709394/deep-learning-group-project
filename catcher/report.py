"""
Teaching Analysis Reporting & Evaluation
========================================

This module is responsible for the final stage of the Teaching Analysis 
Pipeline. It aggregates data from the visual and auditory analysis modules 
to generate structured reports, perform semantic alignment checks, and 
provide pedagogical feedback via LLM integration.

Key Features:
    * Multi-format report generation (PDF, Word, and JSON) with styling.
    * DeepSeek LLM integration for automated teaching verdict and feedback.
    * Semantic alignment analysis between board-written text and spoken content.
    * Robust tokenization and mathematical symbol handling for STEM content.
    * Fallback mechanisms and error logging for API-dependent evaluations.

Dependencies:
    * reportlab: For dynamic PDF generation and custom layout management.
    * python-docx: For programmatic creation of Microsoft Word reports.
    * requests: For communicating with the DeepSeek/LLM REST API.
    * transformers: For local semantic scoring and perplexity evaluation.

Author: [Lai Tsz Yeung/Group J]
Date: 2026
License: MIT
"""
from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Set

from requests import Session, RequestException

# PDF Generation Imports
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

# Word Generation Imports (Requires: pip install python-docx)
try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants / Thresholds
# ---------------------------------------------------------------------------

DEFAULT_BASE_URL = "https://api.deepseek.com/v1"
DEFAULT_MODEL = "deepseek-chat"
DEFAULT_API_KEY_ENV = "DEEPSEEK_API_KEY"
DEFAULT_SBERT = "sentence-transformers/all-mpnet-base-v2"
KEYWORD_OVERLAP_FOR_PARTIAL = 0.15
MATH_SYMBOLS = set("=+-*/??÷^()[]{}<>?????????????·.,:;%")

# ---------------------------------------------------------------------------
# Tokenization / Overlap
# ---------------------------------------------------------------------------

def _tokenize_mixed(text: str) -> Set[str]:
    if not text or not text.strip():
        return set()
    words = re.findall(r"[a-zA-Z]+", text.lower())
    chars = re.findall(r"[\u4e00-\u9fff]", text)
    digits = re.findall(r"\d+", text)
    return set(words) | set(chars) | set(digits)

def _keyword_overlap_rate(a: str, b: str) -> float:
    sa, sb = _tokenize_mixed(a), _tokenize_mixed(b)
    if not sa and not sb:
        return 1.0
    inter = len(sa & sb)
    union = len(sa | sb)
    return float(inter / union) if union else 0.0

# ---------------------------------------------------------------------------
# Heuristic alignment (SBERT + keyword overlap)
# ---------------------------------------------------------------------------

class SemanticAligner:
    def __init__(self, model_name: str = DEFAULT_SBERT) -> None:
        try:
            from sentence_transformers import SentenceTransformer, util
            self._model = SentenceTransformer(model_name)
            self._util = util
        except ImportError:
            logger.warning("sentence_transformers not installed. SBERT alignment will fail.")
            self._model = None
            self._util = None

    def similarity(self, a: str, b: str) -> float:
        if not self._model:
            raise RuntimeError("sentence_transformers is required for SemanticAligner")
        
        # Encode a and b independently or grab their specific indices
        emb = self._model.encode([a, b], convert_to_tensor=True, show_progress_bar=False)
        
        # Calculate similarity specifically between embedding 0 (a) and embedding 1 (b)
        sim = self._util.cos_sim(emb[0], emb[1]).item()
        
        return float(sim)

def _get_alignment_model(model_name: str) -> SemanticAligner:
    return SemanticAligner(model_name)

def _judge_alignment(
    semantic_sim: float,
    keyword_overlap: float,
    high_sim: float = 0.72,
    partial_sim: float = 0.45,
    keyword_high: float = 0.35,
) -> str:
    strong_meaning = semantic_sim >= high_sim
    enough_shared_words = keyword_overlap >= keyword_high
    if strong_meaning and enough_shared_words:
        return "highly_aligned"

    meaning_close_enough = semantic_sim >= partial_sim
    some_shared_words = keyword_overlap >= KEYWORD_OVERLAP_FOR_PARTIAL
    if meaning_close_enough or some_shared_words:
        return "partially_related"

    return "content_mismatch"

def compare_board_and_speech(
    board_text: str,
    speech_text: str,
    *,
    model_name: str = DEFAULT_SBERT,
    high_sim: float = 0.72,
    partial_sim: float = 0.45,
    keyword_high: float = 0.35,
    aligner: Optional[SemanticAligner] = None,
) -> Dict[str, Any]:
    print("compare_board_and_speech")
    sbert = aligner or SemanticAligner(model_name)
    try:
        cosine_similarity = sbert.similarity(board_text, speech_text)
    except Exception:
        logger.exception("SBERT similarity failed")
        raise

    token_overlap = _keyword_overlap_rate(board_text, speech_text)
    verdict = _judge_alignment(
        cosine_similarity,
        token_overlap,
        high_sim=high_sim,
        partial_sim=partial_sim,
        keyword_high=keyword_high,
    )
    return {
        "semantic_similarity": round(cosine_similarity, 4),
        "keyword_overlap_rate": round(token_overlap, 4),
        "verdict": verdict,
    }

# ---------------------------------------------------------------------------
# DeepSeek Plumbing & Filtering
# ---------------------------------------------------------------------------

def _filter_noise_board_lines(lines: List[str], min_chars: int, min_letters: int) -> List[str]:
    filtered: list[str] = []
    for line in lines:
        text = (line or "").strip()
        if len(text) < max(1, min_chars):
            continue
        if sum(ch.isalpha() or ch in MATH_SYMBOLS for ch in text) < max(1, min_letters):
            continue
        filtered.append(text)
    return filtered

def _build_filter_board_lines_messages(lines: List[str], speech_text: str) -> List[Dict[str, str]]:
    system_prompt = (
        "You clean noisy OCR lines from a classroom whiteboard, projector slide, or similar. "
        "Remove: UI chrome, button labels, stray single words that are not educational content, "
        "watermarks, timestamps, window titles, duplicate near-empty fragments, and obvious OCR garbage. "
        "Keep: instructional text, headings, bullet content, formula words, and phrases that belong to the lesson. "
        "Prefer lines that are primarily English/Latin script; drop obvious non-Latin script noise when it is not lesson content. "
        "speech_text is optional context only (what the teacher said); use it to prefer lines that match the topic, "
        "but do not invent lines. Return JSON only."
    )
    numbered = "\n".join(f"{i}: {line}" for i, line in enumerate(lines))
    user_prompt = (
        "Task:\n"
        "1. Review each numbered OCR line below.\n"
        "2. Decide which line indices to KEEP (educational / slide content).\n"
        "3. When uncertain, prefer KEEP over DROP.\n\n"
        "Output requirements:\n"
        '- Return exactly one JSON object with these fields only:\n'
        '  "kept_indices": array of integers ??? 0-based indices referring ONLY to the numbered list below\n'
        '  "reason": one short sentence explaining the main noise you removed\n'
        "Preserve the order of indices as they appear in the list (ascending).\n\n"
        f"speech_text (context, may be empty):\n{(speech_text or '').strip() or '[EMPTY]'}\n\n"
        f"numbered OCR lines:\n{numbered}\n"
    )
    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

def _build_messages(board_text: str, speech_text: str) -> List[Dict[str, str]]:
    # Redirects to the alignment messages builder as requested by duplicate legacy code
    return _build_alignment_messages(board_text, speech_text)

def _build_alignment_messages(board_text: str, speech_text: str) -> List[Dict[str, str]]:
    system_prompt = (
        "You are an evaluator for classroom teaching alignment. "
        "You compare board_text and speech_text, then judge whether the speech stays focused on the board content "
        "or is clearly off-topic. "
        "Return JSON only. Do not include markdown, prose outside JSON, or extra keys."
    )
    user_prompt = (
        "Task:\n"
        "1. Compare the two inputs: board_text and speech_text.\n"
        "2. Judge whether the speech is centered on the board content.\n"
        "3. Judge whether the speech is clearly off-topic.\n"
        "4. Be conservative when information is missing, empty, or too short. Do not guess details.\n\n"
        "Output requirements:\n"
        '- Return exactly one JSON object with these fields only:\n'
        '  "overall_relevance": one of ["highly_relevant", "partially_relevant", "weakly_relevant", "off_topic"]\n'
        '  "score": a number from 0 to 100\n'
        '  "reason": a short explanation in one or two sentences\n'
        '  "evidence": an array of 2 to 5 short strings\n'
        '  "summary" : a string of a summary of the video (board_text and speech_text came from the same video)\n'
        "- If the texts are empty or insufficient, give a cautious low-confidence style judgment rather than making things up.\n\n"
        f"board_text:\n{board_text.strip() or '[EMPTY]'}\n\n"
        f"speech_text:\n{speech_text.strip() or '[EMPTY]'}"
    )
    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

def _normalize_endpoint(base_url: str) -> str:
    cleaned = (base_url or DEFAULT_BASE_URL).strip().rstrip("/")
    if cleaned.endswith("/chat/completions"):
        return cleaned
    return f"{cleaned}/chat/completions"

def _call_deepseek_chat_completion(
    url: str,
    api_key: str,
    model: str,
    timeout_sec: float,
    messages: List[Dict[str, str]],
) -> Dict[str, Any]:
    payload = {
        "model": model,
        "messages": messages,
        "stream": False,
    }
    body = json.dumps(payload).encode("utf-8")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }

    session = Session()
    try:
        resp = session.post(url, data=body, headers=headers, timeout=float(timeout_sec))
        resp.raise_for_status()
    except RequestException as e:
        raise RuntimeError(f"DeepSeek request failed: {e}") from e

    try:
        response_json = resp.json()
    except ValueError as e:
        raise RuntimeError(f"DeepSeek returned invalid JSON response: {resp.text[:500]}") from e

    if not isinstance(response_json, dict):
        raise RuntimeError("DeepSeek response root is not a JSON object")
    return response_json

def _extract_message_content(response_json: Dict[str, Any]) -> str:
    choices = response_json.get("choices")
    if not isinstance(choices, list) or not choices:
        raise RuntimeError("DeepSeek response missing choices")
    message = choices[0].get("message") if isinstance(choices[0], dict) else None
    if not isinstance(message, dict):
        raise RuntimeError("DeepSeek response missing message")
    content = message.get("content")
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        parts: List[str] = []
        for item in content:
            if isinstance(item, dict) and isinstance(item.get("text"), str):
                parts.append(item["text"])
        joined = "".join(parts).strip()
        if joined:
            return joined
    raise RuntimeError("DeepSeek response missing text content")

def _extract_json_object(text: str) -> Dict[str, Any]:
    cleaned = (text or "").strip()
    if not cleaned:
        raise ValueError("Empty DeepSeek content")
    if cleaned.startswith("```"):
        lines = cleaned.splitlines()
        if len(lines) >= 3:
            cleaned = "\n".join(lines[1:-1]).strip()
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("No JSON object found in DeepSeek content")
    snippet = cleaned[start : end + 1]
    parsed = json.loads(snippet)
    if not isinstance(parsed, dict):
        raise ValueError("DeepSeek content JSON is not an object")
    return parsed

def _deepseek_communication(
    messages: List[Dict[str, str]],
    config: Any,
    env: Any,
) -> Optional[Tuple[str, Dict[str, Any]]]:
    url = _normalize_endpoint(getattr(config, "deepseek_base_url", DEFAULT_BASE_URL))
    api_key = getattr(env, "deepseek", os.getenv(DEFAULT_API_KEY_ENV, "sk-6a304dfb56ec43958e10ed366b8b961f")) or "sk-6a304dfb56ec43958e10ed366b8b961f"
    print(api_key)
    if not api_key:
        logger.error("DeepSeek API key not provided")
        return None

    try:
        response_json = _call_deepseek_chat_completion(
            url=url,
            api_key=api_key,
            model=getattr(config, "deepseek_model", DEFAULT_MODEL),
            timeout_sec=float(getattr(config, "deepseek_timeout_sec", 20.0)),
            messages=messages,
        )
        content = _extract_message_content(response_json)
        parsed = _extract_json_object(content)
    except Exception as e:
        logger.exception("DeepSeek communication failed: %s", e)
        return None
    return content, parsed

def deepseek_filter_board_lines(
    board_lines: List[str],
    entire_speech: str,
    config: Any,
    env: Any,
) -> Tuple[List[str], Optional[Dict[str, Any]]]:
    messages = _build_filter_board_lines_messages(board_lines, entire_speech)
    result = _deepseek_communication(messages, config, env)
    if result is None:
        return board_lines, None
    _, keep_dict = result

    kept_indices = keep_dict.get("kept_indices", []) if isinstance(keep_dict, dict) else []
    if not isinstance(kept_indices, list):
        kept_indices = []
    kept_set = {int(i) for i in kept_indices if 0 <= int(i) < len(board_lines)}

    filtered = [line for idx, line in enumerate(board_lines) if idx in kept_set]
    return filtered, keep_dict

def deepseek_alignment_evaluate(
    board_text: str,
    speech_text: str,
    config: Any,
    env: Any,
) -> Optional[Dict[str, Any]]:
    print("asking for summary")
    messages = _build_alignment_messages(board_text, speech_text)
    result = _deepseek_communication(messages, config, env)
    if result is None:
        return None
    _, parsed = result
    if not isinstance(parsed, dict):
        print("no parsed")
        return None
    print(parsed)
    return parsed

# ---------------------------------------------------------------------------
# Wrapped "Main" Executions (Prevents crash on import)
# ---------------------------------------------------------------------------

def run_filter_main(board_lines_primary: List[str], entire_speech: str, config: Any, env: Any) -> Any:
    # Encapsulates "the first main function"
    messages = _build_filter_board_lines_messages(board_lines_primary, entire_speech)
    return _deepseek_communication(messages, config, env)

def run_alignment_main(board_lines_primary: str, entire_speech: str, config: Any, env: Any) -> Any:
    # Encapsulates "the 3rd main function"
    messages = _build_messages(board_lines_primary, entire_speech)
    return _deepseek_communication(messages, config, env)

# ---------------------------------------------------------------------------
# Export Functions (PDF / Word / JSON)
# ---------------------------------------------------------------------------

def _register_cjk_font() -> str:
    # A basic fallback font provider to ensure canvas doesn't crash on CJK chars
    return "Helvetica"

def _build_teaching_feedback_pdf(
    output_path: str, *, board_lines: List[str], clarity: Dict,
    dsalignment: Dict, alignment: Dict, speech_text: str, module_errors: Optional[Dict] = None,
) -> str:
    font = _register_cjk_font()
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(path), pagesize=A4)
    w, h = A4
    y = h - 50
    line_h = 16

    def draw_line(text: str, indent: int = 0) -> None:
        nonlocal y
        if y < 60:
            c.showPage()
            y = h - 50
        c.setFont(font, 11)
        safe = text.encode("latin-1", "replace").decode("latin-1") if font == "Helvetica" else text
        try:
            c.drawString(50 + indent, y, safe[:120])
        except Exception:
            c.drawString(50 + indent, y, safe.encode("ascii", "replace").decode("ascii")[:120])
        y -= line_h

    draw_line("Classroom blackboard analytics - teaching feedback report")
    y -= 8
    draw_line(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    y -= 8

    # Section 1: recognized lines from the board image
    draw_line("1. Board OCR")
    if board_lines:
        for t in board_lines:
            draw_line(f"- {t}", indent=10)
    else:
        draw_line("(no text recognized)", indent=10)
    y -= 8

    # Section 2: legibility score and short advice
    draw_line("2. Handwriting clarity")
    if clarity:
        draw_line(f"Level: {clarity.get('clarity', '')}  Score: {clarity.get('score', '')}", indent=10)
        draw_line(f"Suggestion: {clarity.get('suggestion', '')}", indent=10)
        draw_line(
            f"Laplacian variance: {clarity.get('laplacian_variance', '')}  "
            f"Stroke width variance: {clarity.get('stroke_width_variance', '')}",
            indent=10,
        )
    y -= 8

    # Section 3: raw transcript
    draw_line("3. Speech summary (Whisper)")
    st = speech_text or "(missing or transcription failed)"
    for chunk in range(0, len(st), 90):
        draw_line(st[chunk : chunk + 90], indent=10)
    y -= 8

    # Section 4: do board and speech agree?
    draw_line("4. Board vs speech alignment")
    if alignment:
        draw_line(f"Semantic similarity: {alignment.get('semantic_similarity')}", indent=10)
        draw_line(f"Keyword overlap (Jaccard): {alignment.get('keyword_overlap_rate')}", indent=10)
        draw_line(f"Verdict: {alignment.get('verdict')}", indent=10)
    else:
        draw_line("(skipped or unavailable)", indent=10)

    draw_line("5. Summary")
    if dsalignment:
        st = dsalignment.get('summary') or "(missing or summarize failed)"
        for chunk in range(0, len(st), 90):
            draw_line(st[chunk : chunk + 90], indent=10)
    else:
        draw_line("(skipped or unavailable)", indent=10)


    if module_errors:
        y -= 8
        draw_line("6. Steps that reported an error")
        for step_name, message in module_errors.items():
            if message:
                draw_line(f"{step_name}: {message}", indent=10)

    c.save()
    return str(path.resolve())


def _build_teaching_feedback_word(
    output_path: str, *, board_lines: List[str], clarity: Dict,
    dsalignment: Dict, alignment: Dict, speech_text: str, module_errors: Optional[Dict] = None,
) -> str:
    if docx is None:
        raise RuntimeError("python-docx is not installed. Run 'pip install python-docx'.")
        
    doc = docx.Document()
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.add_heading('Classroom Blackboard Analytics', 0)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    
    doc.add_heading('1. Board OCR', level=1)
    if board_lines:
        for t in board_lines:
            doc.add_paragraph(t, style='List Bullet')
    else:
        doc.add_paragraph("(no text recognized)")
        
    doc.add_heading('2. Handwriting clarity', level=1)
    if clarity:
        doc.add_paragraph(f"Level: {clarity.get('clarity', '')} | Score: {clarity.get('score', '')}")
        doc.add_paragraph(f"Suggestion: {clarity.get('suggestion', '')}")
        
    doc.add_heading('3. Speech summary (Whisper)', level=1)
    doc.add_paragraph(speech_text or "(missing or transcription failed)")
    
    doc.add_heading('4. Board vs speech alignment', level=1)
    if alignment:
        doc.add_paragraph(f"Semantic similarity: {alignment.get('semantic_similarity')}")
        doc.add_paragraph(f"Keyword overlap (Jaccard): {alignment.get('keyword_overlap_rate')}")
        doc.add_paragraph(f"Verdict: {alignment.get('verdict')}")
    else:
        doc.add_paragraph("(skipped or unavailable)")

    doc.add_heading('5. Summary', level=1)
    if dsalignment:
        st = dsalignment.get('summary') or "(missing or summarize failed)"
        doc.add_paragraph(st)

    if module_errors:
        doc.add_heading('6. Steps that reported an error', level=1)
        for step_name, message in module_errors.items():
            if message:
                doc.add_paragraph(f"{step_name}: {message}", style='List Bullet')
                
    doc.save(str(path))
    return str(path.resolve())

def tmp_save_pdf(output_path: str, payload: Dict[str, Any]) -> Dict[str, Optional[str]]:
    lines = payload.get("board_lines") or []
    clarity_block = payload.get("clarity") or {}
    align_block = payload.get("alignment")
    dsalign_block = payload.get("deepseek_alignment_verdict")
    spoken = payload.get("speech_text") or ""
    failures = payload.get("module_errors")

    out = {"pdf_path": None, "error": None}
    try:
        out["pdf_path"] = _build_teaching_feedback_pdf(
            output_path,
            board_lines=lines,
            clarity=clarity_block,
            alignment=align_block,
            dsalignment=dsalign_block,
            speech_text=spoken,
            module_errors=failures,
        )
    except Exception as e:
        logger.exception("Failed to build PDF")
        out["error"] = str(e)
    return out

def tmp_save_word(output_path: str, payload: Dict[str, Any]) -> Dict[str, Optional[str]]:
    lines = payload.get("board_lines") or []
    clarity_block = payload.get("clarity") or {}
    align_block = payload.get("alignment")
    dsalign_block = payload.get("deepseek_alignment_verdict")
    spoken = payload.get("speech_text") or ""
    failures = payload.get("module_errors")

    out = {"word_path": None, "error": None}
    try:
        out["word_path"] = _build_teaching_feedback_word(
            output_path,
            board_lines=lines,
            clarity=clarity_block,
            alignment=align_block,
            dsalignment=dsalign_block,
            speech_text=spoken,
            module_errors=failures,
        )
    except Exception as e:
        logger.exception("Failed to build Word")
        out["error"] = str(e)
    return out

def tmp_save_json(output_path: str, payload: Dict[str, Any]) -> Dict[str, Optional[str]]:
    out = {"json_path": None, "error": None}
    path = Path(output_path)
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=4, ensure_ascii=False)
        out["json_path"] = str(path.resolve())
    except Exception as e:
        logger.exception("Failed to save JSON")
        out["error"] = str(e)
    return out